import os
from docx import Document
from docx.shared import Pt, Inches
from tqdm import tqdm
from io import StringIO
import sys
import re
from docx.styles.style import WD_STYLE_TYPE


def execute_python_code(file_path):
    stdout_backup = sys.stdout
    stderr_backup = sys.stderr

    sys.stdout = StringIO()
    sys.stderr = StringIO()

    try:
        with open(file_path, "r") as file:
            code = file.read()
        exec(code, {})

        stdout_value = sys.stdout.getvalue()
        stderr_value = sys.stderr.getvalue()
    except Exception as e:
        stdout_value = ""
        stderr_value = f"Произошла ошибка: {str(e)}"
    finally:
        stdin_value = sys.stdin.getvalue()

        sys.stdout = stdout_backup
        sys.stderr = stderr_backup

    return stdout_value, stderr_value, stdin_value


def add_code_listing_to_doc(doc, code_lines):
    for i, line in enumerate(code_lines):
        p = doc.add_paragraph(f"{i + 1:4} {line.strip()}", style="Code")
        p.style.font.name = "Courier New"
        p.style.font.size = Pt(10)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(0)  # Убираем межстрочный интервал


def sort_files(files):
    def get_file_key(file_name):
        match = re.search(r"(\d+)", file_name)
        number = int(match.group(1)) if match else float("inf")
        if file_name.startswith("task"):
            return (0, number)
        elif file_name.startswith("extra"):
            return (1, number)
        return (2, number)

    return sorted(files, key=get_file_key)


def create_styles(doc):
    styles = doc.styles

    heading1 = styles.add_style("Heading1", WD_STYLE_TYPE.PARAGRAPH)
    heading1.font.name = "Times New Roman"
    heading1.font.size = Pt(18)
    heading1.font.bold = True

    heading2 = styles.add_style("Heading2", WD_STYLE_TYPE.PARAGRAPH)
    heading2.font.name = "Times New Roman"
    heading2.font.size = Pt(14)
    heading2.font.bold = True

    code_style = styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = "Courier New"
    code_style.font.size = Pt(10)
    code_style.paragraph_format.left_indent = Inches(0.5)
    code_style.paragraph_format.space_after = Pt(0)


def generate_report(folder_path, output_file):
    doc = Document()
    create_styles(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(0.6)

    # Заголовок отчета
    p = doc.add_paragraph("Отчет о проделанной работе", style="Heading1")
    p.alignment = 1  # Выравнивание по центру

    python_files = [f for f in os.listdir(folder_path) if f.endswith(".py")]
    python_files = sort_files(python_files)

    for file_name in tqdm(python_files, desc="Обработка файлов"):
        file_path = os.path.join(folder_path, file_name)

        # Заголовок для каждого файла
        p = doc.add_paragraph(f"Отчет по файлу: {file_name}", style="Heading2")

        with open(file_path, "r") as file:
            code_lines = file.readlines()

        # Заголовок для кода
        p = doc.add_paragraph("Исходный код", style="Heading2")

        add_code_listing_to_doc(doc, code_lines)

        stdout, stderr, stdin = execute_python_code(file_path)

        # Заголовок для результатов выполнения
        p = doc.add_paragraph("Результаты выполнения", style="Heading2")

        if stdin:
            p = doc.add_paragraph("Ввод:", style="Normal")
            doc.add_paragraph(stdin, style="Normal")
        if stdout:
            p = doc.add_paragraph("Вывод программы:", style="Normal")
            doc.add_paragraph(stdout, style="Normal")
        if stderr:
            p = doc.add_paragraph("Ошибки:", style="Normal")
            doc.add_paragraph(stderr, style="Normal")

        # Добавляем разрыв страницы перед отчетом по каждому файлу
        doc.add_page_break()
    doc.save(output_file)
    print(f"Отчет сохранен как {output_file}")


if __name__ == "__main__":
    generate_report("src/prac1", "work_report.docx")
